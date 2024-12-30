import spacy
import re
from typing import Dict, List, Optional
from datetime import datetime
from pydantic import BaseModel

class Education(BaseModel):
    degree: str
    school: str
    graduation_date: Optional[str]
    gpa: Optional[float]

class WorkExperience(BaseModel):
    company: str
    title: str
    start_date: str
    end_date: Optional[str]
    description: List[str]
    skills: List[str]

class ResumeSection(BaseModel):
    education: List[Education]
    work_experience: List[WorkExperience]
    skills: List[str]
    contact: Dict[str, str]
    projects: List[Dict[str, str]]

class ResumeParser:
    def __init__(self):
        self.nlp = spacy.load("en_core_web_lg")
        
        # Common section headers
        self.section_headers = {
            'education': ['education', 'academic background', 'academic history', 'academic qualification'],
            'experience': ['experience', 'work experience', 'employment history', 'work history', 'professional experience'],
            'skills': ['skills', 'technical skills', 'core competencies', 'competencies'],
            'projects': ['projects', 'personal projects', 'professional projects'],
            'contact': ['contact', 'contact information', 'personal information']
        }
        
        # Common skill keywords
        self.technical_skills = set([
            'python', 'java', 'javascript', 'react', 'node', 'sql', 'aws', 'docker',
            'kubernetes', 'machine learning', 'ai', 'data science', 'devops', 'cloud',
            'git', 'agile', 'scrum', 'ci/cd', 'rest api', 'microservices'
        ])

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text based on file type with improved formatting preservation."""
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(file_content)
            elif file_type == 'docx':
                return self._extract_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")

    def _extract_from_pdf(self, content: bytes) -> str:
        import pdfplumber  # Better text extraction than PyPDF2
        
        text_content = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                text_content.append(page.extract_text(x_tolerance=3))
        
        return "\n".join(text_content)

    def _extract_from_docx(self, content: bytes) -> str:
        from docx import Document
        
        doc = Document(BytesIO(content))
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)

    def parse_sections(self, text: str) -> ResumeSection:
        """Parse resume into structured sections."""
        doc = self.nlp(text)
        sections = self._identify_sections(text)
        
        return ResumeSection(
            education=self._parse_education(sections.get('education', '')),
            work_experience=self._parse_experience(sections.get('experience', '')),
            skills=self._parse_skills(sections.get('skills', '')),
            contact=self._parse_contact(sections.get('contact', '')),
            projects=self._parse_projects(sections.get('projects', ''))
        )

    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify different sections in the resume."""
        lines = text.split('\n')
        sections = {}
        current_section = None
        current_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line is a section header
            section_found = False
            for section, headers in self.section_headers.items():
                if any(header in line.lower() for header in headers):
                    if current_section:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = section
                    current_content = []
                    section_found = True
                    break

            if not section_found and current_section:
                current_content.append(line)

        # Add the last section
        if current_section:
            sections[current_section] = '\n'.join(current_content)

        return sections

    def _parse_education(self, text: str) -> List[Education]:
        """Extract education information."""
        education_list = []
        if not text:
            return education_list

        # Split into different education entries
        entries = re.split(r'\n(?=[A-Z])', text)
        
        for entry in entries:
            degree = None
            school = None
            graduation_date = None
            gpa = None

            # Extract degree
            degree_patterns = [
                r'(Bachelor|Master|PhD|B\.S\.|M\.S\.|B\.A\.|M\.A\.|Ph\.D)\.?\s+(?:of|in)?\s+([^,\n]+)',
                r'([^,\n]+?)\s+Degree'
            ]
            for pattern in degree_patterns:
                match = re.search(pattern, entry, re.IGNORECASE)
                if match:
                    degree = match.group(0)
                    break

            # Extract school
            school_match = re.search(r'(University|College|Institute|School)\s+of\s+[^,\n]+', entry)
            if school_match:
                school = school_match.group(0)

            # Extract graduation date
            date_match = re.search(r'(May|June|December|January|Spring|Fall)\s+\d{4}', entry)
            if date_match:
                graduation_date = date_match.group(0)

            # Extract GPA
            gpa_match = re.search(r'GPA:\s*(\d+\.\d+)', entry)
            if gpa_match:
                gpa = float(gpa_match.group(1))

            if degree or school:
                education_list.append(Education(
                    degree=degree or "Unknown Degree",
                    school=school or "Unknown School",
                    graduation_date=graduation_date,
                    gpa=gpa
                ))

        return education_list

    def _parse_experience(self, text: str) -> List[WorkExperience]:
        """Extract work experience information."""
        experience_list = []
        if not text:
            return experience_list

        # Split into different job entries
        entries = re.split(r'\n(?=[A-Z][^,\n]+(?:Inc\.|LLC|Ltd\.|\||\-))', text)
        
        for entry in entries:
            # Extract company
            company_match = re.search(r'^([^,|\n-]+)', entry)
            company = company_match.group(1).strip() if company_match else "Unknown Company"

            # Extract title
            title_match = re.search(r'(?:^|\n)([^,|\n-]+(?:Engineer|Developer|Manager|Director|Analyst|Designer)[^,\n]+)', entry)
            title = title_match.group(1).strip() if title_match else "Unknown Title"

            # Extract dates
            dates_match = re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s\-]+\d{4}[\s\-]+(?:to[\s\-]+)?(Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s\-]+\d{4})?', entry)
            start_date = "Unknown"
            end_date = "Present"
            if dates_match:
                dates = dates_match.group(0).split('to')
                start_date = dates[0].strip()
                if len(dates) > 1:
                    end_date = dates[1].strip()

            # Extract description and skills
            description_lines = []
            skills = set()
            
            lines = entry.split('\n')
            for line in lines:
                line = line.strip()
                if line and not any(x in line.lower() for x in ['company', 'position', 'present']):
                    description_lines.append(line)
                    # Extract skills from description
                    words = set(word.lower() for word in line.split())
                    skills.update(words.intersection(self.technical_skills))

            experience_list.append(WorkExperience(
                company=company,
                title=title,
                start_date=start_date,
                end_date=end_date,
                description=description_lines,
                skills=list(skills)
            ))

        return experience_list

    def _parse_skills(self, text: str) -> List[str]:
        """Extract skills from the skills section."""
        if not text:
            return []

        # Split skills by common delimiters
        skills = set()
        skill_text = re.sub(r'[\n•●■\-\|,]', ' ', text)
        
        # Extract mentioned skills
        doc = self.nlp(skill_text)
        for token in doc:
            if token.text.lower() in self.technical_skills:
                skills.add(token.text.lower())

        return list(skills)

    def extract_keywords(self, text: str) -> Dict[str, List[str]]:
        """Extract important keywords by category."""
        doc = self.nlp(text)
        
        keywords = {
            'technical_skills': [],
            'soft_skills': [],
            'tools': [],
            'certifications': [],
            'industries': []
        }
        
        # Add keyword extraction logic here
        for token in doc:
            if token.text.lower() in self.technical_skills:
                keywords['technical_skills'].append(token.text)
            
            # Add more keyword categorization logic
        
        return keywords

    def calculate_match_score(self, resume_text: str, job_description: str) -> Dict[str, any]:
        """Calculate how well the resume matches a job description."""
        resume_doc = self.nlp(resume_text)
        job_doc = self.nlp(job_description)
        
        # Calculate similarity score
        similarity_score = resume_doc.similarity(job_doc)
        
        # Extract keywords from both
        resume_keywords = self.extract_keywords(resume_text)
        job_keywords = self.extract_keywords(job_description)
        
        # Calculate matching and missing keywords
        matching_keywords = set(resume_keywords['technical_skills']).intersection(set(job_keywords['technical_skills']))
        missing_keywords = set(job_keywords['technical_skills']) - set(resume_keywords['technical_skills'])
        
        return {
            'similarity_score': similarity_score,
            'matching_keywords': list(matching_keywords),
            'missing_keywords': list(missing_keywords)
        }